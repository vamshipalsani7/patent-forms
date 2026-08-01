"""Document → text, dispatched by file type.

This is the ONLY stage of the extraction pipeline that varies by file type.
Everything downstream — classifier, profile_builder, extractors — consumes
``[(page_number, text), …]`` and neither knows nor cares whether that text came
from a PDF, a Word document, or a plain text file. That invariant is the whole
point of this module: it lets the product accept the formats patent professionals
actually hold (a specification is as often a .docx as a .pdf) without a second
copy of the extraction logic per format.

Supported, in descending fidelity:
  .pdf   → the existing PDFReader (pdfplumber), page by page.
  .docx  → python-docx, paragraphs and table cells, treated as a single page.
  .txt   → decoded text, treated as a single page.
  .doc   → best-effort only. The legacy binary Word format (OLE compound file)
           has no dependable pure-Python reader; a faithful extraction needs
           LibreOffice or antiword, which are not carried here. We attempt a
           conservative salvage and, if it yields nothing usable, report the
           file as unreadable rather than feeding the classifier binary noise.

A reader never raises for an unreadable document. It returns empty pages, and the
caller records an honest "couldn't read this file" status — the same graceful
degradation the PDF path already had for image-only scans.

Page semantics: PDFs keep real page numbers. DOCX/TXT/DOC have no page structure
in their source, so they report as a single page (1). A Fact's provenance still
answers "which document", and "page 1 of a Word file" is truthful — the value did
come from the one page there is.
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Union

from extractor.pdf_reader import PDFReader


class UnsupportedDocumentError(ValueError):
    """The file extension is not one this reader handles."""


# Extensions we accept for upload. Kept here so the API and the reader cannot
# disagree about what "supported" means.
SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".doc", ".txt"})

# Content types the browser may send for the supported extensions. Advisory
# only — extension is the authority, because browsers report Word/text types
# inconsistently across platforms.
SUPPORTED_CONTENT_TYPES = frozenset({
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
    "application/msword",  # .doc
    "text/plain",
    "application/octet-stream",  # some browsers send this for local files
    "",
})


class DocumentReader:
    """Reads raw text from an uploaded document, dispatching on its extension."""

    def __init__(self, pdf_reader: PDFReader | None = None) -> None:
        self._pdf_reader = pdf_reader or PDFReader()

    def read_pages(
        self, file_path: Union[str, Path], filename: str | None = None
    ) -> list[tuple[int, str]]:
        """Return ``[(page_number, text), …]`` for the document.

        Args:
            file_path: Path to the file on disk.
            filename:  Original filename, used to pick the reader when the path
                       on disk is a content-addressed name without an extension.
                       Falls back to ``file_path``'s own suffix.

        Empty pages are preserved so page numbers stay truthful. An unreadable
        document returns ``[]`` — no exception — so extraction degrades to "no
        facts" rather than crashing the upload.
        """
        suffix = self._suffix(file_path, filename)
        if suffix == ".pdf":
            return self._pdf_reader.read_pages(file_path)
        if suffix == ".docx":
            return self._read_docx(file_path)
        if suffix == ".txt":
            return self._read_txt(file_path)
        if suffix == ".doc":
            return self._read_doc_best_effort(file_path)
        raise UnsupportedDocumentError(
            f"Unsupported document type '{suffix}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )

    def read(self, file_path: Union[str, Path], filename: str | None = None) -> str:
        """Full text of the document, pages joined — parallels PDFReader.read()."""
        return "\n\n".join(t for _, t in self.read_pages(file_path, filename) if t)

    # ------------------------------------------------------------------ helpers

    @staticmethod
    def _suffix(file_path: Union[str, Path], filename: str | None) -> str:
        """The extension that decides the reader, preferring the original name.

        Uploads are stored under a content id like ``doc_ms3d46fw_bfxavr.pdf``;
        that suffix is trustworthy today, but taking the original filename's
        suffix when given keeps this correct even if storage naming changes.
        """
        name = filename if filename else str(file_path)
        return Path(name).suffix.lower()

    # -------------------------------------------------------------------- .docx

    @staticmethod
    def _read_docx(file_path: Union[str, Path]) -> list[tuple[int, str]]:
        """Extract paragraphs and table cells from a .docx as one page.

        python-docx is imported lazily so the rest of the pipeline (and the whole
        test suite) still loads if the optional dependency is absent — a missing
        library reports as an unreadable document, not an import crash at startup.
        """
        try:
            import docx  # type: ignore
        except ImportError:
            return []

        try:
            document = docx.Document(str(file_path))
        except Exception:  # noqa: BLE001 — corrupt or not a real .docx
            return []

        lines: list[str] = [p.text for p in document.paragraphs]
        # Party details on IPO-style documents very often sit in tables; a
        # paragraphs-only read would miss the applicant/inventor blocks entirely.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("\t".join(c for c in cells if c))

        text = "\n".join(line for line in lines if line is not None)
        return [(1, text)] if text.strip() else []

    # --------------------------------------------------------------------- .txt

    @staticmethod
    def _read_txt(file_path: Union[str, Path]) -> list[tuple[int, str]]:
        """Decode a plain-text file, tolerant of encoding.

        UTF-8 first (with BOM handling), then a Latin-1 fallback that cannot
        raise — the goal is to get readable text to the classifier, not to be
        strict about encoding.
        """
        data = Path(file_path).read_bytes()
        for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                text = data.decode(encoding)
                return [(1, text)] if text.strip() else []
            except UnicodeDecodeError:
                continue
        return []

    # --------------------------------------------------------------------- .doc

    @staticmethod
    def _read_doc_best_effort(file_path: Union[str, Path]) -> list[tuple[int, str]]:
        """Attempt to salvage text from a legacy binary .doc; may return [].

        There is deliberately no heavy dependency here. Two cheap attempts:

          1. Some files with a .doc name are actually .docx (a zip). Detect that
             by magic bytes and read them properly.
          2. Otherwise, extract runs of printable characters from the OLE stream.
             This is genuinely lossy — it recovers visible words but not
             structure — so it is a fallback, not a supported path.

        If neither yields usable text, return [] and let the caller report the
        file as unreadable. That honesty matters more than a garbled half-read
        landing on a statutory form. See module docstring.
        """
        path = Path(file_path)

        # (1) Mislabelled .docx (PK zip signature).
        try:
            if zipfile.is_zipfile(path):
                return DocumentReader._read_docx(path)
        except Exception:  # noqa: BLE001
            pass

        # (2) Printable-run salvage from the raw bytes.
        try:
            data = path.read_bytes()
        except OSError:
            return []

        text = DocumentReader._salvage_printable_runs(data)
        return [(1, text)] if text.strip() else []

    @staticmethod
    def _salvage_printable_runs(data: bytes, min_run: int = 4) -> str:
        """Pull runs of readable Latin-script text out of binary Word bytes.

        Word 97-2003 stores document text as UTF-16LE, so ASCII letters appear
        as ``<char>\\x00``. We read the little-endian UTF-16 view and keep runs,
        at least ``min_run`` long, of characters in the Latin range only.

        The Latin-only restriction is deliberate. Arbitrary binary decoded as
        UTF-16LE readily yields *printable* high-plane codepoints (CJK and the
        like) — accepting those would turn structural noise into "text" and rob
        the caller of its honest "couldn't read this" signal. IPO documents are
        Latin script, so a run outside that range is far likelier to be noise
        than content. This recovers party names and titles often enough to be
        worth attempting, and returns nothing when there is nothing real to find.
        """
        try:
            decoded = data.decode("utf-16-le", errors="ignore")
        except Exception:  # noqa: BLE001
            return ""

        def is_text(char: str) -> bool:
            code = ord(char)
            if char in "\t\n\r":
                return True
            # ASCII printable, or Latin-1 / Latin Extended letters and marks.
            return (0x20 <= code <= 0x7E) or (0xA0 <= code <= 0x24F)

        runs: list[str] = []
        current: list[str] = []
        for char in decoded:
            if is_text(char):
                current.append(char)
            else:
                if len(current) >= min_run:
                    runs.append("".join(current))
                current = []
        if len(current) >= min_run:
            runs.append("".join(current))

        # A salvage that is almost all punctuation/whitespace is noise, not text.
        joined = "\n".join(runs)
        letters = sum(1 for c in joined if c.isalpha())
        return joined if letters >= min_run else ""
