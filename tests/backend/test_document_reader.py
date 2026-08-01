"""DocumentReader — the one pipeline stage that varies by file type.

The guarantee under test is that every supported format arrives downstream as
the same [(page, text), …] shape, so the classifier and extractors never learn
what a file was. The reader must also fail soft: an unreadable document yields
no pages, never an exception, because a crash on upload is worse than a document
that produced nothing.
"""

from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

import context  # noqa: F401  — sets sys.path

from extractor.document_reader import (
    SUPPORTED_EXTENSIONS,
    DocumentReader,
    UnsupportedDocumentError,
)


class TestPlainText(unittest.TestCase):
    def setUp(self):
        self.reader = DocumentReader()
        self.tmp = Path(tempfile.mkdtemp())

    def _write(self, name, data, encoding="utf-8"):
        path = self.tmp / name
        path.write_bytes(data.encode(encoding) if isinstance(data, str) else data)
        return path

    def test_reads_utf8_text_as_a_single_page(self):
        path = self._write("spec.txt", "COMPLETE SPECIFICATION\nTitle: A Widget")
        pages = self.reader.read_pages(path)
        self.assertEqual(1, len(pages))
        self.assertEqual(1, pages[0][0])
        self.assertIn("COMPLETE SPECIFICATION", pages[0][1])

    def test_reads_windows_encoded_text(self):
        path = self._write("cp1252.txt", "Café Ünïcode".encode("cp1252"))
        self.assertIn("Café", self.reader.read(path))

    def test_empty_text_file_yields_no_pages(self):
        path = self._write("blank.txt", "   \n  ")
        self.assertEqual([], self.reader.read_pages(path))


class TestDocx(unittest.TestCase):
    def setUp(self):
        self.reader = DocumentReader()
        self.tmp = Path(tempfile.mkdtemp())

    def _docx(self, paragraphs, table_rows=None):
        import docx

        document = docx.Document()
        for para in paragraphs:
            document.add_paragraph(para)
        if table_rows:
            table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
            for r, row in enumerate(table_rows):
                for c, val in enumerate(row):
                    table.rows[r].cells[c].text = val
        path = self.tmp / "doc.docx"
        document.save(path)
        return path

    def test_reads_paragraphs(self):
        path = self._docx(["FORM 26", "Authorisation of a Patent Agent"])
        text = self.reader.read(path)
        self.assertIn("FORM 26", text)
        self.assertIn("Authorisation", text)

    def test_reads_table_cells(self):
        """Party details on IPO documents commonly sit in tables — a
        paragraphs-only reader would miss the applicant block."""
        path = self._docx(["Applicant details:"], table_rows=[["Name", "Acme Ltd"]])
        text = self.reader.read(path)
        self.assertIn("Acme Ltd", text)

    def test_a_docx_flows_through_classification_identically_to_a_pdf(self):
        from extractor.classifier import DocumentClassifier
        from extractor.profile_builder import ProfileBuilder
        from models.patent_profile import DocumentType

        path = self._docx([
            "FORM 26",
            "FORM FOR AUTHORISATION OF A PATENT AGENT",
            "I/We, Acme Ltd, do hereby authorise Shri RAJESH KUMAR, IN/PA-1234, to act.",
        ])
        pages = self.reader.read_pages(path, "poa.docx")
        full_text = "\n\n".join(t for _, t in pages)

        self.assertEqual(
            DocumentType.FORM26_AUTHORISATION,
            DocumentClassifier().classify(full_text),
        )
        extract = ProfileBuilder().build(
            file_path=path, document_id="d1", original_filename="poa.docx",
            document_type=DocumentType.FORM26_AUTHORISATION, page_texts=pages,
        )
        self.assertIn("RAJESH KUMAR", [f.value for f in extract.facts])


class TestDoc(unittest.TestCase):
    def setUp(self):
        self.reader = DocumentReader()
        self.tmp = Path(tempfile.mkdtemp())

    def test_a_docx_mislabelled_as_doc_is_still_read(self):
        import docx

        document = docx.Document()
        document.add_paragraph("PATENT CERTIFICATE")
        path = self.tmp / "mislabelled.doc"
        document.save(path)  # writes a real .docx (zip) under a .doc name
        self.assertTrue(zipfile.is_zipfile(path))
        self.assertIn("PATENT CERTIFICATE", self.reader.read(path, "mislabelled.doc"))

    def test_binary_word_text_is_salvaged_when_present(self):
        """Word 97-2003 stores text as UTF-16LE; the salvage recovers visible
        words even without a full binary parser."""
        payload = "PATENT CERTIFICATE Patentee ACME INNOVATIONS".encode("utf-16-le")
        blob = b"\xd0\xcf\x11\xe0" + b"\x00" * 32 + payload + b"\x00" * 16
        path = self.tmp / "legacy.doc"
        path.write_bytes(blob)
        self.assertIn("PATENT CERTIFICATE", self.reader.read(path, "legacy.doc"))

    def test_unsalvageable_doc_yields_no_pages_rather_than_raising(self):
        path = self.tmp / "noise.doc"
        path.write_bytes(bytes(range(256)) * 4)  # binary noise, no readable runs
        self.assertEqual([], self.reader.read_pages(path, "noise.doc"))


class TestDispatchAndFailure(unittest.TestCase):
    def setUp(self):
        self.reader = DocumentReader()
        self.tmp = Path(tempfile.mkdtemp())

    def test_unsupported_extension_raises(self):
        path = self.tmp / "sheet.xlsx"
        path.write_bytes(b"anything")
        with self.assertRaises(UnsupportedDocumentError):
            self.reader.read_pages(path, "sheet.xlsx")

    def test_extension_is_taken_from_the_original_filename_when_given(self):
        """Uploads are stored under a content id; the original name decides the
        reader so a future storage-naming change cannot misroute."""
        path = self.tmp / "doc_abc123.bin"
        path.write_text("plain notes", encoding="utf-8")
        self.assertIn("plain notes", self.reader.read(path, "notes.txt"))

    def test_supported_extensions_are_the_advertised_set(self):
        self.assertEqual({".pdf", ".docx", ".doc", ".txt"}, set(SUPPORTED_EXTENSIONS))


if __name__ == "__main__":
    unittest.main()
