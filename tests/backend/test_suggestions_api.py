"""Suggestion generation and the content store.

These call the endpoint functions directly rather than over HTTP: the project
has no httpx installed (which FastAPI's TestClient requires), and the routing
layer is thin enough that the value is in the pipeline behind it. Every test
redirects the content store at a temp directory so the real backend/uploads/
is never touched.
"""

from __future__ import annotations

import asyncio
import io
import json
import tempfile
import unittest
from pathlib import Path

import context  # noqa: F401  — sets sys.path

import storage.content_store as content_store_module
from fastapi import HTTPException, UploadFile

import app as app_module
from models.document_extract import DocumentExtract
from models.fact import Fact


def _fact(key, value, source_type="form1", confidence=0.8, document_id="doc_1"):
    return Fact(
        key=key, value=value, document_id=document_id, source_type=source_type,
        page=1, confidence=confidence, method="anchor", extractor_version="form1@1",
    )


WORKSPACE = "default"


def _extract(document_id, *facts, source_type="form1", workspace_id=WORKSPACE):
    return DocumentExtract(
        document_id=document_id, source_type=source_type,
        workspace_id=workspace_id,
        original_filename=f"{document_id}.pdf", page_count=1,
        facts=list(facts), extractor_version="form1@1",
    )


class IsolatedStoreTestCase(unittest.TestCase):
    """Points the module-level content store at a throwaway directory."""

    def setUp(self):
        self._tmp = tempfile.TemporaryDirectory()
        self._original_dir = content_store_module._UPLOADS_DIR
        content_store_module._UPLOADS_DIR = Path(self._tmp.name)
        self.store = content_store_module.ContentStore()
        self._original_store = app_module._store
        app_module._store = self.store

    def tearDown(self):
        app_module._store = self._original_store
        content_store_module._UPLOADS_DIR = self._original_dir
        self._tmp.cleanup()

    def upload(self, document_id, pdf_path, filename=None,
               content_type="application/pdf", workspace_id=WORKSPACE):
        data = Path(pdf_path).read_bytes()
        upload = UploadFile(
            file=io.BytesIO(data),
            filename=filename or Path(pdf_path).name,
            headers={"content-type": content_type},
        )
        return asyncio.run(app_module.extract_document(
            file=upload, document_id=document_id, workspace_id=workspace_id,
        ))


class TestSuggestionGeneration(IsolatedStoreTestCase):
    def test_returns_the_documented_response_shape(self):
        self.store.save_extract(_extract("d1", _fact("applicant.name", "Acme Ltd")))
        response = app_module.get_suggestions("form_03", WORKSPACE)

        self.assertEqual("form_03", response["form_id"])
        self.assertIsInstance(response["suggestions"], dict)

    def test_each_suggestion_carries_value_and_provenance(self):
        self.store.save_extract(_extract("d1", _fact("applicant.name", "Acme Ltd")))
        suggestions = app_module.get_suggestions("form_03", WORKSPACE)["suggestions"]

        self.assertIn("applicant_declaration.applicant_names", suggestions)
        entry = suggestions["applicant_declaration.applicant_names"]
        self.assertIn("value", entry)
        for field in ("document_id", "source_type", "page", "confidence",
                      "method", "extractor_version"):
            self.assertIn(field, entry["fact"], f"provenance is missing '{field}'")

    def test_no_documents_means_no_suggestions(self):
        response = app_module.get_suggestions("form_03", WORKSPACE)
        self.assertEqual({}, response["suggestions"])

    def test_a_document_with_no_facts_produces_no_suggestions(self):
        self.store.save_extract(_extract("d1", source_type="unknown"))
        self.assertEqual({}, app_module.get_suggestions("form_03", WORKSPACE)["suggestions"])

    def test_unknown_form_id_raises_404(self):
        with self.assertRaises(HTTPException) as caught:
            app_module.get_suggestions("form_99_does_not_exist", WORKSPACE)
        self.assertEqual(404, caught.exception.status_code)

    def test_form_id_is_not_used_to_traverse_the_filesystem(self):
        with self.assertRaises(HTTPException) as caught:
            app_module.get_suggestions("../../../etc/passwd", WORKSPACE)
        self.assertEqual(404, caught.exception.status_code)


class TestContentStore(IsolatedStoreTestCase):
    def test_saves_and_returns_the_pdf_path(self):
        path = self.store.save_pdf(WORKSPACE, "d1", b"%PDF-1.4 fake bytes")
        self.assertTrue(path.exists())
        self.assertEqual(path, self.store.get_pdf_path(WORKSPACE, "d1"))

    def test_returns_none_for_an_unknown_pdf(self):
        self.assertIsNone(self.store.get_pdf_path(WORKSPACE, "never_uploaded"))

    def test_round_trips_an_extract(self):
        self.store.save_extract(_extract("d1", _fact("applicant.name", "Acme Ltd")))
        loaded = self.store.get_extract(WORKSPACE, "d1")

        self.assertIsNotNone(loaded)
        self.assertEqual("Acme Ltd", loaded.facts[0].value)
        self.assertEqual("form1", loaded.source_type)

    def test_re_extraction_replaces_the_previous_record(self):
        self.store.save_extract(_extract("d1", _fact("applicant.name", "Old Name")))
        self.store.save_extract(_extract("d1", _fact("applicant.name", "New Name")))

        self.assertEqual(1, len(self.store.extracts_for_workspace(WORKSPACE)))
        self.assertEqual("New Name", self.store.get_extract(WORKSPACE, "d1").facts[0].value)

    def test_extracts_survive_a_restart(self):
        """Extracts are rebuildable, but must not be lost on process restart."""
        self.store.save_extract(_extract("d1", _fact("applicant.name", "Acme Ltd")))
        reopened = content_store_module.ContentStore()

        self.assertEqual(1, len(reopened.extracts_for_workspace(WORKSPACE)))
        self.assertEqual("Acme Ltd", reopened.get_extract(WORKSPACE, "d1").facts[0].value)

    def test_a_corrupt_sidecar_is_skipped_rather_than_fatal(self):
        self.store.save_extract(_extract("d1", _fact("applicant.name", "Acme Ltd")))
        (Path(self._tmp.name) / WORKSPACE / "d2.extract.json").write_text("{ corrupt", encoding="utf-8")

        reopened = content_store_module.ContentStore()
        self.assertEqual(["d1"], [e.document_id for e in reopened.extracts_for_workspace(WORKSPACE)])


class TestExtractEndpointGuards(IsolatedStoreTestCase):
    def test_rejects_an_unsupported_file_type(self):
        """The type is decided by extension; an unsupported one is refused."""
        upload = UploadFile(
            file=io.BytesIO(b"<xml/>"), filename="drawing.rtf",
            headers={"content-type": "application/rtf"},
        )
        with self.assertRaises(HTTPException) as caught:
            asyncio.run(app_module.extract_document(
                file=upload, document_id="d1", workspace_id=WORKSPACE,
            ))
        self.assertEqual(400, caught.exception.status_code)

    def test_accepts_a_supported_non_pdf_type(self):
        """A .txt source document is now accepted, not rejected — content type
        is advisory, the extension is the authority."""
        upload = UploadFile(
            file=io.BytesIO(b"COMPLETE SPECIFICATION\nApplication No. : 202211012345"),
            filename="notes.txt", headers={"content-type": "text/plain"},
        )
        result = asyncio.run(app_module.extract_document(
            file=upload, document_id="d_txt", workspace_id=WORKSPACE,
        ))
        self.assertEqual("d_txt", result["document_id"])
        self.assertNotIn("error", result)

    def test_content_type_is_advisory_not_authoritative(self):
        """A real PDF whose browser-reported content type is wrong (or blank)
        must still be processed by its extension."""
        upload = UploadFile(
            file=io.BytesIO(b"%PDF-1.4 not really a pdf"), filename="spec.pdf",
            headers={"content-type": "application/octet-stream"},
        )
        result = asyncio.run(app_module.extract_document(
            file=upload, document_id="d_pdf", workspace_id=WORKSPACE,
        ))
        # It is accepted (no 400); extraction may still find nothing, which is fine.
        self.assertEqual("d_pdf", result["document_id"])

    def test_rejects_an_empty_upload(self):
        upload = UploadFile(
            file=io.BytesIO(b""), filename="empty.pdf",
            headers={"content-type": "application/pdf"},
        )
        with self.assertRaises(HTTPException) as caught:
            asyncio.run(app_module.extract_document(
                file=upload, document_id="d1", workspace_id=WORKSPACE,
            ))
        self.assertEqual(400, caught.exception.status_code)

    def test_an_unreadable_pdf_degrades_instead_of_raising_500(self):
        """The form must stay usable when extraction fails."""
        upload = UploadFile(
            file=io.BytesIO(b"%PDF-1.4 this is not really a pdf"),
            filename="broken.pdf", headers={"content-type": "application/pdf"},
        )
        result = asyncio.run(app_module.extract_document(
            file=upload, document_id="d_broken", workspace_id=WORKSPACE,
        ))

        self.assertIn("error", result)
        self.assertEqual([], result["facts"])
        self.assertEqual("d_broken", result["document_id"])


class TestSuggestionsOverrides(IsolatedStoreTestCase):
    """The suggestions endpoint carries the user's Workspace decisions through
    to the generated form via the optional `overrides` JSON param."""

    def _title_suggestion(self, result):
        for path, s in result["suggestions"].items():
            if path.endswith("invention_title") or path.endswith(".invention.title"):
                return s
        return None

    def test_a_resolved_conflict_reaches_the_form(self):
        self.store.save_extract(_extract(
            "spec", _fact("invention.title", "Spec Title", source_type="form2_specification"),
            source_type="form2_specification",
        ))
        self.store.save_extract(_extract(
            "cert", _fact("invention.title", "Certificate Title", source_type="patent_certificate"),
            source_type="patent_certificate",
        ))

        overrides = json.dumps({"invention.title": "Certificate Title"})
        result = app_module.get_suggestions("form_01", workspace_id=WORKSPACE, overrides=overrides)
        s = self._title_suggestion(result)
        self.assertIsNotNone(s)
        self.assertEqual("Certificate Title", s["value"])

    def test_a_typed_value_reaches_the_form_as_user_provided(self):
        self.store.save_extract(_extract(
            "spec", _fact("invention.title", "A Title", source_type="form2_specification"),
            source_type="form2_specification",
        ))
        overrides = json.dumps({"applicant.nationality": "Indian"})
        result = app_module.get_suggestions("form_01", workspace_id=WORKSPACE, overrides=overrides)

        provided = [
            s for s in result["suggestions"].values()
            if s["value"] == "Indian" and s["fact"]["source_type"] == "user"
        ]
        self.assertTrue(provided, "the typed nationality did not reach the form")

    def test_malformed_overrides_are_ignored_not_fatal(self):
        self.store.save_extract(_extract(
            "spec", _fact("invention.title", "A Title", source_type="form2_specification"),
            source_type="form2_specification",
        ))
        # Not JSON at all — the form must still generate from the extractions.
        result = app_module.get_suggestions("form_01", workspace_id=WORKSPACE, overrides="{not json")
        self.assertIn("suggestions", result)
        self.assertTrue(result["suggestions"])

    def test_no_overrides_param_is_the_prior_behaviour(self):
        self.store.save_extract(_extract(
            "spec", _fact("invention.title", "A Title", source_type="form2_specification"),
            source_type="form2_specification",
        ))
        result = app_module.get_suggestions("form_01", workspace_id=WORKSPACE)
        self.assertIn("suggestions", result)


class TestWorkspaceEndpoint(IsolatedStoreTestCase):
    """GET /api/workspace/{id} — the Patent Workspace's data source."""

    def test_summarises_the_workspaces_documents_and_facts(self):
        self.store.save_extract(_extract(
            "spec", _fact("invention.title", "A Widget"),
            _fact("applicant.name", "Acme Ltd"), source_type="form2_specification",
        ))
        summary = app_module.get_workspace(WORKSPACE)

        self.assertEqual(WORKSPACE, summary["workspace_id"])
        self.assertEqual(1, summary["stats"]["document_count"])
        section_ids = [s["id"] for s in summary["sections"]]
        self.assertIn("title", section_ids)
        self.assertIn("applicants", section_ids)

    def test_is_scoped_to_one_workspace(self):
        self.store.save_extract(_extract("a", _fact("applicant.name", "Company A"),
                                         workspace_id="patent_a"))
        self.store.save_extract(_extract("b", _fact("applicant.name", "Company B"),
                                         workspace_id="patent_b"))

        summary_a = app_module.get_workspace("patent_a")
        self.assertEqual(1, summary_a["stats"]["document_count"])
        names = [
            v["value"]
            for s in summary_a["sections"] if s["id"] == "applicants"
            for f in s["fields"] for v in f["values"]
        ]
        self.assertIn("Company A", names)
        self.assertNotIn("Company B", names)

    def test_an_empty_workspace_is_well_formed(self):
        summary = app_module.get_workspace("never_used_ws")
        self.assertEqual([], summary["documents"])
        self.assertEqual([], summary["sections"])
        self.assertTrue(summary["missing"], "core fields should be reported missing")

    def test_a_bad_workspace_id_is_rejected(self):
        with self.assertRaises(HTTPException) as caught:
            app_module.get_workspace("../escape")
        self.assertEqual(400, caught.exception.status_code)

    def test_an_uploaded_docx_reaches_the_summary_with_a_friendly_type(self):
        """End to end: a Word document, uploaded through the real endpoint,
        appears in the workspace with a human type label — no PDF required."""
        import docx

        tmp = Path(self._tmp.name) / "poa.docx"
        document = docx.Document()
        document.add_paragraph("FORM 26")
        document.add_paragraph("FORM FOR AUTHORISATION OF A PATENT AGENT")
        document.add_paragraph(
            "I/We, Acme Ltd, do hereby authorise Shri RAJESH KUMAR, IN/PA-1234, to act."
        )
        document.save(tmp)

        result = self.upload("d_docx", tmp, filename="poa.docx", content_type="")
        self.assertEqual("form26_authorisation", result["source_type"])
        self.assertEqual("Form 26 (Authorisation of Agent)", result["source_type_label"])

        summary = app_module.get_workspace(WORKSPACE)
        doc = summary["documents"][0]
        self.assertEqual("poa.docx", doc["filename"])
        self.assertEqual("Form 26 (Authorisation of Agent)", doc["document_type"])
        self.assertEqual("extracted", doc["status"])


class TestWorkspaceScopingBoundary(IsolatedStoreTestCase):
    """Cross-matter isolation is now enforced in the backend, not just the UI.

    Merging a Form 1 for Patent A with a certificate for Patent B produces a
    confidently wrong profile — the worst failure mode, because it looks fine.
    Full coverage lives in test_workspace_isolation.py; these two keep the
    guarantee asserted at the point it used to be missing.
    """

    def test_documents_are_pooled_only_within_their_own_workspace(self):
        self.store.save_extract(
            _extract("doc_a", _fact("applicant.name", "Company A"), workspace_id="patent_a")
        )
        self.store.save_extract(
            _extract("doc_b", _fact("applicant.name", "Company B"), workspace_id="patent_b")
        )

        self.assertEqual(1, len(self.store.extracts_for_workspace("patent_a")))
        self.assertEqual(1, len(self.store.extracts_for_workspace("patent_b")))
        self.assertEqual(0, len(self.store.extracts_for_workspace(WORKSPACE)))

    def test_a_profile_only_ever_sees_the_extracts_it_is_given(self):
        """The isolation primitive: scoping is a matter of what you pass in."""
        from models.patent_profile import PatentProfile

        extract_a = _extract("patent_a", _fact("applicant.name", "Company A"))
        extract_b = _extract("patent_b", _fact("applicant.name", "Company B"))

        profile_a = PatentProfile(extracts=[extract_a])
        profile_b = PatentProfile(extracts=[extract_b])

        self.assertEqual(["Company A"], [f.value for f in profile_a.get_facts("applicant.name")])
        self.assertEqual(["Company B"], [f.value for f in profile_b.get_facts("applicant.name")])


if __name__ == "__main__":
    unittest.main()
